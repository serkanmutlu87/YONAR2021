#!/usr/bin/env python
#-*-coding:utf-8-*-
import pandas as pd
import docx
import time
from docx2pdf import convert

df = pd.read_excel(r"veriler.xlsx", sheet_name="Sayfa1")

print("Çalışma Başlıyor...")

for it in range(len(df["İsim Soyisim"])):
    print("\n\n", df["Bildiri İsmi"][it], " için oluşturuluyor.")
    doc = docx.Document("taslak.docx")
    doc.paragraphs[6].text = df["İsim Soyisim"][it]
    doc.paragraphs[7].text = df["Bildiri İsmi"][it]
    doc.save(df["İsim Soyisim"][it] + " - " + df["Bildiri İsmi"][it][:30] + "..." +".docx")
    convert(df["İsim Soyisim"][it] + " - " + df["Bildiri İsmi"][it][:30] + "..." +".docx")

time.pause(100)